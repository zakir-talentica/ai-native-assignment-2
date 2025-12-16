import os
from typing import List, Tuple
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from ..config import settings


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx."""
    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_md(file_path: str) -> str:
    """Read markdown file as plain text."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def extract_text(file_path: str) -> str:
    """Extract text based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.md':
        return extract_text_from_md(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """Chunk text using LangChain's RecursiveCharacterTextSplitter."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    chunks = text_splitter.split_text(text)
    return chunks


def generate_embeddings_and_index(
    chunks: List[str],
    metadata: List[dict],
    index_path: str
) -> None:
    """Generate embeddings and add to FAISS index."""
    embeddings = OpenAIEmbeddings(openai_api_key=settings.openai_api_key)
    
    # Create Document objects with metadata
    documents = [
        Document(page_content=chunk, metadata=meta)
        for chunk, meta in zip(chunks, metadata)
    ]
    
    # Load existing or create new index
    if os.path.exists(f"{index_path}/index.faiss"):
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        vectorstore.add_documents(documents)
    else:
        os.makedirs(index_path, exist_ok=True)
        vectorstore = FAISS.from_documents(documents, embeddings)
    
    # Save the updated index
    vectorstore.save_local(index_path)


def load_faiss_index(index_path: str) -> FAISS:
    """Load FAISS index from disk."""
    embeddings = OpenAIEmbeddings(openai_api_key=settings.openai_api_key)
    return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)


def process_document(file_path: str, document_id: str) -> None:
    """Process a document: extract, chunk, embed, and index."""
    # Extract text
    text = extract_text(file_path)
    
    # Chunk text
    chunks = chunk_text(text)
    
    # Create metadata for each chunk
    filename = os.path.basename(file_path)
    metadata = [
        {
            "document_id": document_id,
            "filename": filename,
            "chunk_id": f"{document_id}_chunk_{i}",
            "chunk_index": i
        }
        for i in range(len(chunks))
    ]
    
    # Generate embeddings and add to index
    generate_embeddings_and_index(chunks, metadata, settings.faiss_index_path)

